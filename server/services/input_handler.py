# src/input_handler.py
import os
import mimetypes
from typing import Union
from docx import Document
from PyPDF2 import PdfReader

def detect_input_type(input_data: Union[str, bytes]) -> str:
    """
    判断输入数据是纯文本还是文档文件。
    - 如果是文件路径，根据文件扩展名判断 'text', 'docx', 'pdf'
    - 如果是字符串且无文件扩展名，则归为 'text'
    """
    if isinstance(input_data, str) and os.path.isfile(input_data):
        mime, _ = mimetypes.guess_type(input_data)
        if input_data.lower().endswith('.txt') or mime == 'text/plain':
            return 'text'
        elif input_data.lower().endswith('.docx'):
            return 'docx'
        elif input_data.lower().endswith('.pdf'):
            return 'pdf'
        else:
            return 'unknown'
    else:
        # 当成原始文本处理
        return 'text'
def load_content(input_data: Union[str, bytes]) -> str:
    """
    根据输入类型加载文本内容：
    - text: 原样返回
    - docx: 读取所有段落文本
    - pdf: 读取所有页文本
    """
    itype = detect_input_type(input_data)
    if itype == 'text':
        return input_data if isinstance(input_data, str) else input_data.decode('utf-8')
    elif itype == 'docx':
        doc = Document(input_data)
        return "\n".join([para.text for para in doc.paragraphs])
    elif itype == 'pdf':
        reader = PdfReader(input_data)
        pages = [page.extract_text() for page in reader.pages]
        return "\n".join(pages)
    else:
        raise ValueError(f"无法处理的文件类型：{itype}")

def parse_doc_structure(doc_id, filename):
    path = f"user_input/{doc_id}_{filename}"
    ext = os.path.splitext(filename)[-1].lower()
    structure = []
    title_keywords = ['1.', '一、', '引言']
    if ext == ".docx":
        doc = Document(path)
        paras = [para for para in doc.paragraphs if para.text.strip()]
        if paras:
            first_text = paras[0].text.strip()
            if not any(kw in first_text for kw in title_keywords):
                # 首句为标题
                structure.append({"id": "title0", "text": first_text, "type": "title", "level": 1})
                start_idx = 1
            else:
                start_idx = 0
            for i, para in enumerate(paras[start_idx:], start=start_idx):
                if para.style.name == "Heading 1":
                    structure.append({"id": f"title{i}", "text": para.text.strip(), "type": "title", "level": 1})
                elif para.style.name == "Heading 2":
                    structure.append({"id": f"subtitle{i}", "text": para.text.strip(), "type": "subtitle", "level": 2})
                else:
                    structure.append({"id": f"para{i}", "text": para.text.strip(), "type": "paragraph"})
    elif ext == ".txt":
        with open(path, encoding="utf-8") as f:
            lines = [line.strip() for line in f if line.strip()]
        if lines:
            first_text = lines[0]
            if not any(kw in first_text for kw in title_keywords):
                structure.append({"id": "title0", "text": first_text, "type": "title", "level": 1})
                start_idx = 1
            else:
                start_idx = 0
            for i, line in enumerate(lines[start_idx:], start=start_idx):
                structure.append({"id": f"para{i}", "text": line, "type": "paragraph"})
    return structure